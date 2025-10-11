from flask import Flask, render_template, request, jsonify, send_file
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from datetime import datetime
import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

# Global variables
driver = None
navigation_data = []

# Configure Gemini AI
genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
model = genai.GenerativeModel('gemini-pro')

def init_driver():
    """Initialize Chrome WebDriver with options - MANUAL VERSION"""
    try:
        chrome_options = Options()
        chrome_options.add_argument('--start-maximized')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])
        
        # Option 1: Try automatic first
        try:
            from webdriver_manager.chrome import ChromeDriverManager
            driver_path = ChromeDriverManager().install()
            service = Service(executable_path=driver_path)
            driver = webdriver.Chrome(service=service, options=chrome_options)
            print("✅ ChromeDriver loaded automatically")
            return driver
        except:
            # Option 2: Use manual chromedriver.exe from project folder
            chromedriver_path = os.path.join(os.getcwd(), 'chromedriver.exe')
            if os.path.exists(chromedriver_path):
                service = Service(executable_path=chromedriver_path)
                driver = webdriver.Chrome(service=service, options=chrome_options)
                print("✅ ChromeDriver loaded from project folder")
                return driver
            else:
                raise Exception("ChromeDriver not found! Please download chromedriver.exe and place it in the project folder.")
                
    except Exception as e:
        print(f"❌ Error initializing driver: {e}")
        raise

def inject_tracking_script():
    """Inject click tracking script into current page"""
    script = """
    if (!window.navTrackerInjected) {
        window.navTrackerInjected = true;
        
        document.addEventListener('click', function(e) {
            let element = e.target;
            
            // Get the clickable element (could be nested)
            let maxDepth = 5;
            let depth = 0;
            while (element && depth < maxDepth) {
                if (['A', 'BUTTON', 'NAV', 'LI', 'DIV'].includes(element.tagName)) {
                    // Check if it's a navigation element
                    if (element.tagName === 'A' || 
                        element.tagName === 'BUTTON' || 
                        element.closest('nav') || 
                        element.getAttribute('role') === 'button' ||
                        element.getAttribute('role') === 'link' ||
                        element.classList.contains('nav') ||
                        element.classList.contains('menu') ||
                        element.classList.contains('link')) {
                        break;
                    }
                }
                element = element.parentElement;
                depth++;
                if (!element || element === document.body) break;
            }
            
            if (element && element !== document.body) {
                let elementText = element.innerText || 
                                element.textContent || 
                                element.getAttribute('aria-label') || 
                                element.getAttribute('title') ||
                                element.getAttribute('alt') ||
                                element.className ||
                                'Unnamed Element';
                
                let elementData = {
                    text: elementText.trim().substring(0, 100),
                    tag: element.tagName,
                    url: window.location.href,
                    href: element.href || '',
                    className: element.className || '',
                    id: element.id || '',
                    timestamp: new Date().toISOString()
                };
                
                // Store in sessionStorage
                let clicks = JSON.parse(sessionStorage.getItem('navClicks') || '[]');
                clicks.push(elementData);
                sessionStorage.setItem('navClicks', JSON.stringify(clicks));
                
                console.log('✅ Click captured:', elementData.text);
            }
        }, true);
        
        console.log('🎯 Navigation tracker injected on:', window.location.href);
    }
    """
    try:
        driver.execute_script(script)
        return True
    except Exception as e:
        print(f"Error injecting script: {e}")
        return False

def analyze_element_with_ai(element_text, element_tag, url):
    """Use AI to analyze and describe the clicked element"""
    try:
        prompt = f"""Analyze this navigation element and provide a brief, professional description (max 100 chars):
        
Element Text: {element_text}
Element Type: {element_tag}
Page URL: {url}

Describe what this element does or where it leads. Be concise and clear."""
        
        response = model.generate_content(prompt)
        description = response.text.strip()
        return description[:200]  # Limit to 200 chars
    except Exception as e:
        return f"Navigation element: {element_text[:50]}"

def save_navigation_data(element_name, url, description):
    """Save navigation data to the global list"""
    navigation_data.append({
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'element_name': element_name,
        'url': url,
        'description': description
    })

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/start_navigation', methods=['POST'])
def start_navigation():
    """Start the navigation session"""
    try:
        data = request.json
        url = data.get('url', 'https://www.google.com')
        
        global driver
        driver = init_driver()
        driver.get(url)
        
        # Wait a bit for page to load
        driver.implicitly_wait(2)
        
        # Inject tracking script
        inject_tracking_script()
        
        return jsonify({'success': True, 'message': 'Navigation started successfully'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/capture_clicks', methods=['GET'])
def capture_clicks():
    """Capture and process clicked elements"""
    try:
        if not driver:
            return jsonify({'success': False, 'error': 'Driver not initialized'})
        
        # RE-INJECT the tracking script (in case user navigated to new page)
        inject_tracking_script()
        
        # Get clicks from sessionStorage
        clicks = driver.execute_script("return JSON.parse(sessionStorage.getItem('navClicks') || '[]');")
        
        # Get already processed clicks
        processed = driver.execute_script("return JSON.parse(sessionStorage.getItem('processedClicks') || '[]');")
        
        new_entries = []
        newly_processed = []
        
        for click in clicks:
            # Create unique identifier for this click
            click_id = f"{click['text']}_{click['url']}_{click.get('href', '')}_{click.get('timestamp', '')}"
            
            # Skip if already processed
            if click_id in processed:
                continue
                
            element_name = click['text'][:100] if click['text'] else 'Unnamed Element'
            url = click.get('href') or click['url']
            
            # Use AI to analyze the element
            description = analyze_element_with_ai(click['text'], click['tag'], url)
            
            # Save the data
            save_navigation_data(element_name, url, description)
            
            new_entries.append({
                'element_name': element_name,
                'url': url,
                'description': description
            })
            
            newly_processed.append(click_id)
        
        # Update processed clicks list
        if newly_processed:
            driver.execute_script(f"""
                let processed = JSON.parse(sessionStorage.getItem('processedClicks') || '[]');
                processed.push(...{json.dumps(newly_processed)});
                sessionStorage.setItem('processedClicks', JSON.stringify(processed));
            """)
        
        # DON'T clear navClicks - keep accumulating them!
        
        return jsonify({
            'success': True,
            'new_entries': len(new_entries),
            'total_entries': len(navigation_data)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/get_navigation_data', methods=['GET'])
def get_navigation_data():
    """Get all navigation data"""
    return jsonify({'success': True, 'data': navigation_data})

@app.route('/stop_navigation', methods=['POST'])
def stop_navigation():
    """Stop the navigation session"""
    global driver
    try:
        if driver:
            driver.quit()
            driver = None
        return jsonify({'success': True, 'message': 'Navigation stopped'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/clear_data', methods=['POST'])
def clear_data():
    """Clear all navigation data"""
    global navigation_data
    navigation_data = []
    return jsonify({'success': True, 'message': 'Data cleared'})

@app.route('/download_word', methods=['GET'])
def download_word():
    """Generate and download Word document"""
    try:
        if not navigation_data:
            return jsonify({'success': False, 'error': 'No data available'})
        
        filename = generate_word_document(navigation_data)
        return send_file(filename, as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

def generate_word_document(navigation_data):
    """Generate Word document from navigation data"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Navigation History Report', 0)
    title.alignment = 1  # Center alignment
    
    # Metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Entries: {len(navigation_data)}')
    doc.add_paragraph('_' * 80)
    
    # Navigation entries
    for idx, item in enumerate(navigation_data, 1):
        # Entry header
        heading = doc.add_heading(f'Entry #{idx}', level=2)
        
        # Timestamp
        p = doc.add_paragraph()
        p.add_run('Timestamp: ').bold = True
        p.add_run(item['timestamp'])
        
        # Element name
        p = doc.add_paragraph()
        p.add_run('Element Name: ').bold = True
        p.add_run(item['element_name'])
        
        # URL
        p = doc.add_paragraph()
        p.add_run('URL: ').bold = True
        run = p.add_run(item['url'])
        run.font.color.rgb = RGBColor(0, 0, 255)
        
        # AI Description
        p = doc.add_paragraph()
        p.add_run('AI Description: ').bold = True
        p.add_run(item['description'])
        
        doc.add_paragraph('_' * 80)
    
    filename = f'navigation_history_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    return filename

if __name__ == '__main__':
    app.run(debug=True, port=5000)